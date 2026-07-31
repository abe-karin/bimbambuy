import glob
import importlib
import logging
import os
import csv

from langchain_core.documents import Document
from pathlib import Path
from dotenv import load_dotenv
from flask import Flask, jsonify, render_template, request

try:
    from langchain_core.documents import Document as LangchainDocument
except Exception:
    class LangchainDocument:
        def __init__(self, page_content, metadata=None):
            self.page_content = page_content
            self.metadata = metadata or {}

Document = LangchainDocument

# ==========================================
# 1. Configuração da API e Inicialização
# ==========================================
load_dotenv(dotenv_path=os.path.join(os.path.dirname(__file__), ".env"))


recuperador_global = None
agente_global = None


class RecuperadorLocal:
    def __init__(self, documentos):
        self.documentos = documentos

    def invoke(self, pergunta):
        if not self.documentos:
            return []

        pergunta_lower = pergunta.lower()
        palavras_comuns = {"o", "a", "os", "as", "um", "uma", "de", "do", "da", "em", "no", "na", "que", "são", "quais", "como", "para", "com"}
        termos_uteis = [t for t in pergunta_lower.split() if t not in palavras_comuns]
        
        pontuados = []
        for documento in self.documentos:
            texto = documento.page_content.lower()
            score = sum(1 for termo in termos_uteis if termo in texto)
            if score:
                pontuados.append((score, documento))

        pontuados.sort(key=lambda item: item[0], reverse=True)
        return [documento for _, documento in pontuados[:3]]
        
        
def _gerar_resposta_local(pergunta, contexto):
    contexto_limpo = " ".join(contexto.split())
    if not contexto_limpo:
        return "Ainda não há conteúdo suficiente na base de conhecimento para responder."

    contexto_resumido = contexto_limpo[:900]
    if len(contexto_limpo) > 900:
        contexto_resumido += "..."

    pergunta_lower = pergunta.lower()
    contexto_lower = contexto_limpo.lower()

    termos_pagamento = [
        "forma de pagamento",
        "formas de pagamento",
        "meio de pagamento",
        "meios de pagamento",
        "pagamento",
        "pagamentos",
        "pix",
        "boleto",
        "cartão",
        "cartoes",
        "carteira digital",
        "transferência bancária",
        "transferencia bancaria",
    ]

    palavras = [palavra for palavra in pergunta_lower.split() if len(palavra) > 2]
    if any(termo in pergunta_lower for termo in termos_pagamento) and any(termo in contexto_lower for termo in termos_pagamento):
        return (
            "Olá! Na BimBam Buy, aceitamos as seguintes formas de pagamento:\n\n"
            "- Cartões de crédito e débito\n"
            "- Transferência bancária (PIX)\n"
            "- Boletos\n"
            "- Carteiras digitais"
        )

    if any(palavra in contexto_lower for palavra in palavras):
        return f"Resposta baseada no contexto disponível: {contexto_resumido}"

    return f"Não encontrei uma resposta exata no contexto atual. Resumo do conteúdo disponível: {contexto_resumido}"


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".csv", ".docx"}
def carregar_documentos(arquivos):
    documentos = []

    for arquivo in arquivos:
        caminho = Path(arquivo)
        if not caminho.exists() or not caminho.is_file():
            continue

        extensao = caminho.suffix.lower()
        print(f"   Lendo o arquivo: {caminho}")

        if extensao == ".pdf":
            try:
                from langchain_unstructured import UnstructuredLoader
            except Exception:
                documentos.append(
                    Document(
                        page_content=f"[PDF não processado automaticamente: {caminho.name}]",
                        metadata={"source": str(caminho), "file_type": "pdf"},
                    )
                )
                continue

            carregador = UnstructuredLoader(str(caminho), languages=["pt"])
            pedacos_pdf = carregador.load()
            
            # Junta todas as linhas soltas do PDF em um único texto contínuo
            texto_completo = "\n".join(pedaco.page_content for pedaco in pedacos_pdf if pedaco.page_content.strip())
            
            # Salva como um documento único
            documentos.append(
                Document(page_content=texto_completo, metadata={"source": str(caminho), "file_type": "pdf"})
            )
        elif extensao == ".txt":
            texto = caminho.read_text(encoding="utf-8", errors="ignore")
            documentos.append(
                Document(page_content=texto, metadata={"source": str(caminho), "file_type": "txt"})
            )
        elif extensao == ".csv":
            print(f"   Processando tabela CSV nativamente: {caminho}")
            with open(caminho, mode='r', encoding='utf-8-sig') as arquivo_csv:
                # O DictReader lê a primeira linha como cabeçalho automaticamente
                leitor_csv = csv.DictReader(arquivo_csv)
                
                for linha in leitor_csv:
                    # Constrói uma frase amigável para a IA com os dados da linha
                    # Exemplo: "Região/Estado: São Paulo | Prazo: 2 a 5 dias | Valor: R$ 15,00"
                    conteudo = " | ".join(f"{chave}: {valor}" for chave, valor in linha.items() if valor)
                    documentos.append(
                        Document(
                            page_content=conteudo,
                            metadata={"source": str(caminho), "file_type": "csv"}
                        )
                    )
        elif extensao in {".docx"}:
          
            try:
                from docx import Document as DocxDocument
            except ImportError as exc:
                raise ImportError("Para ler arquivos Word, instale python-docx no ambiente.") from exc

            doc = DocxDocument(str(caminho))
            texto = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())

            print(f"SUCESSO! O arquivo {caminho.name} foi lido. Total de caracteres: {len(texto)}")
            documentos.append(
                Document(page_content=texto, metadata={"source": str(caminho), "file_type": "docx"})

            
            )
            continue

    return documentos


def listar_arquivos_documentos(pasta="documentos"):
    if not os.path.isdir(pasta):
        return []

    arquivos = []
    for item in glob.glob(os.path.join(pasta, "*")):
        caminho = Path(item)
        
        # Só processa se for arquivo, tiver a extensão certa e NÃO começar com "~"
        if caminho.is_file() and caminho.suffix.lower() in SUPPORTED_EXTENSIONS:
            if not caminho.name.startswith("~"):
                arquivos.append(str(caminho))

    return sorted(arquivos)


def inicializar_base_conhecimento():
    print("1. Carregando os documentos da base de conhecimento...")

    arquivos = listar_arquivos_documentos()
    if not arquivos:
        raise ValueError("Nenhum documento suportado encontrado na pasta 'documentos'.")

    documentos = carregar_documentos(arquivos)
    if not documentos:
        raise ValueError("Não foi possível extrair texto dos documentos encontrados.")

    print("2. Fatiando o texto em pequenos pedaços (Chunks)...")
    try:
        from langchain_community.vectorstores import Chroma
        from langchain_community.vectorstores.utils import filter_complex_metadata
        from langchain_huggingface import HuggingFaceEmbeddings
        from langchain_text_splitters import RecursiveCharacterTextSplitter

        divisor_texto = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        pedacos = divisor_texto.split_documents(documentos)
        pedacos = filter_complex_metadata(pedacos)

        print("3. Criando Embeddings e salvando no Banco de Dados Vetorial (Chroma)...")
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2")
        banco_vetorial = Chroma.from_documents(documents=pedacos, embedding=embeddings)

        recuperador = banco_vetorial.as_retriever(
        search_type="mmr",
        search_kwargs={"k": 6, "fetch_k": 20}
    )

        return recuperador
    except Exception as exc:
        print(f"  Usando recuperação local por fallback: {exc}")
        return RecuperadorLocal(documentos)


def configurar_agente(recuperador):
    google_api_key = os.getenv("GOOGLE_API_KEY")
    openai_api_key = os.getenv("OPENAI_API_KEY")
    os.environ["GOOGLE_API_KEY"] = google_api_key or ""
    os.environ["OPENAI_API_KEY"] = openai_api_key or ""

    print("4. Configurando o cérebro do Agente (LLM + Prompt)...")

    from langchain_core.prompts import ChatPromptTemplate

    prompt_sistema = (
        "Você é o assistente virtual de atendimento da loja online 'BimBam Buy'.\n"
        "Use APENAS os pedaços de contexto recuperados abaixo para responder à pergunta.\n"
        "Se a resposta não estiver no contexto, diga amigavelmente que não encontrou "
        "essa informação nas políticas da loja e oriente a contatar o suporte.\n"
        "Responda de forma clara, educada e direta.\n\n"
        "DICA: Quando o usuário perguntar sobre 'tratamento de dados', procure no contexto por informações sobre coleta, armazenamento, processamento e utilização de dados (LGPD).\n"
        "CONTEXTO RECUPERADO:\n{context}"
    )

    prompt = ChatPromptTemplate.from_messages([
        ("system", prompt_sistema),
        ("human", "{input}"),
    ])

    provedores = []
    if openai_api_key:
        try:
            langchain_openai = importlib.import_module("langchain_openai")
            ChatOpenAI = langchain_openai.ChatOpenAI
            provedores.append(("openai", ChatOpenAI(model="gpt-4o-mini", temperature=0.3)))
        except Exception as exc:
            print(f"  OpenAI indisponível: {exc}")

    if google_api_key:
        try:
            langchain_google_genai = importlib.import_module("langchain_google_genai")
            ChatGoogleGenerativeAI = langchain_google_genai.ChatGoogleGenerativeAI
            provedores.append(("gemini", ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0.3)))
        except Exception as exc:
            print(f"  Gemini indisponível: {exc}")

    class AgenteRAG:
        def __init__(self, recuperador, prompt, provedores):
            self.recuperador = recuperador
            self.prompt = prompt
            self.provedores = provedores

        def invoke(self, payload):
            pergunta = payload["input"]
            documentos_recuperados = self.recuperador.invoke(pergunta)
            contexto = "\n\n".join(doc.page_content for doc in documentos_recuperados)
            

            print("\n" + "="*50)
            print("CONTEXTO ENTREGUE PARA A IA LER:")
            print(contexto[:1500])
            print("="*50 + "\n")

            mensagens = self.prompt.format_messages(input=pergunta, context=contexto)

            ultimo_erro = None
            for nome, llm in self.provedores:
                try:
                    resposta = llm.invoke(mensagens)
                    return {"answer": resposta.content}
                except Exception as exc:
                    ultimo_erro = exc
                    print(f"  Falha com {nome}: {exc}")

            return {"answer": _gerar_resposta_local(pergunta, contexto)}

    if provedores:
        return AgenteRAG(recuperador, prompt, provedores)

    return AgenteRAG(recuperador, prompt, [])


def get_agente():
    global recuperador_global, agente_global

    if agente_global is None:
        print("-" * 50)
        print("Iniciando o sistema RAG da BimBam Buy...")
        recuperador_global = inicializar_base_conhecimento()
        agente_global = configurar_agente(recuperador_global)
        print("Pronto! O Agente Inteligente está online.")
        print("-" * 50)

    return agente_global


# ==========================================
# 2. Configuração do Flask (Rotas Web)
# ==========================================
app = Flask(__name__)

# Configura o arquivo onde as mensagens serão salvas
logging.basicConfig(
    filename='historico_chat.log', # Nome do arquivo que será criado
    level=logging.INFO,
    format='%(asctime)s - %(message)s', # Formato: Data/Hora - Mensagem
    datefmt='%Y-%m-%d %H:%M:%S',
    encoding='utf-8'
)

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/chat', methods=['POST'])
def chat():
    dados = request.get_json()
    pergunta_usuario = dados.get('mensagem')

    if not pergunta_usuario:
        return jsonify({"erro": "Nenhuma mensagem fornecida"}), 400

    try:
        agente = get_agente()
        resposta_agente = agente.invoke({"input": pergunta_usuario})
        resposta_texto = resposta_agente["answer"]

        logging.info(f"PERGUNTA: {pergunta_usuario}")
        logging.info(f"RESPOSTA: {resposta_texto}\n{'-'*30}")

        return jsonify({"resposta": resposta_texto})

    except Exception as e:
        logging.error(f"ERRO: {str(e)}")
        mensagem_segura = "Desculpe, ocorreu um erro ao processar sua resposta. Tente novamente em instantes."
        return jsonify({"resposta": mensagem_segura}), 200


if __name__ == '__main__':
    app.run(debug=True, use_reloader=False)